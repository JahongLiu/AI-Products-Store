import openai

import json
import os
from docx import Document
import docx
import docx.oxml.ns as ns
from docxtpl import DocxTemplate
from docx2pdf import convert
from pypdf import PdfMerger
import PyPDF2
from app.ai_book_generation.templates.template import Template
from app.ai_book_generation.content.ebook import Ebook
from app.ai_book_generation.content.gpt_systems import GptSystems
from app.gpt_wrapper.gpt_wrapper import GptWrapper
from app.ai_book_generation.util.retry import retry
import platform
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import re


class EBookContentGenerator:
    def __init__(self, templates_file, output_directory, id, topic="", target_audience="", recipient_email=""):
        self.GPT_WRAPPER = GptWrapper()
        self.templates_file = templates_file
        self.cover_location = f"app/ai_book_generation/temp/cover-{id}.docx"
        self.cover_pdf_location = f"app/ai_book_generation/temp/cover-{id}.pdf"
        self.cover_photo_location = (
            f"app/ai_book_generation/temp/cover_photo-{id}.jpg"
        )
        self.content_pdf_location_temp = (  # book without cover
            f"app/ai_book_generation/temp/book-{id}-temp.pdf"
        )
        self.content_pdf_location = (  # book without cover
            f"app/ai_book_generation/temp/book-{id}.pdf"
        )
        self.output_pdf_location = os.path.join(
            output_directory, f"final-{id}-{recipient_email}-{topic}-{target_audience}.pdf"
        )  # book with cover

    @retry(max_retries=3)
    def generate_cover(
        self, cover_template, title, topic, target_audience, output_file, preview
    ):
        doc = DocxTemplate(cover_template)

        if preview:
            self.cover_photo_location = f"app/ai_book_generation/templates/covers/preview_photo.png"
        else:
            self.generate_cover_photo(
                title, topic, target_audience, self.cover_photo_location
            )
        imagen = InlineImage(
            doc, self.cover_photo_location, width=Mm(120)
        )  # width is in millimetres
        context = {"title": title, "subtext": "NJ Publishing", "image": imagen}
        doc.render(context)
        doc.save(self.cover_location)
        convert(self.cover_location, output_file)

    @retry(max_retries=3)
    def generate_cover_photo(self, title, topic, target_audience, img_output):
        convo_id = self.GPT_WRAPPER.start_convo(GptSystems.AUTHOR_SYSTEM)
        cover_prompt = (
            f'We have a ebook with the title {title}. It is about "{topic}".'
            f' Our reader is:  "{target_audience}". Write me a very brief and'
            " matter-of-fact description of a photo that would be on the"
            " cover of the book. Do not reference the cover or photo in your"
            ' answer. For example, if the title was "How to lose weight for'
            ' middle aged women", a reasonable response would be "a middle'
            ' age woman exercising"'
        )
        print(cover_prompt)
        dalle_prompt = self.GPT_WRAPPER.msg_in_convo(convo_id, cover_prompt)
        print(dalle_prompt)
        img_data = self.GPT_WRAPPER.generate_photo(dalle_prompt)
        with open(img_output, "wb") as handler:
            handler.write(img_data)

    def merge_pdfs(self, input_files, output_file):
        merger = PdfMerger()

        for pdf in input_files:
            merger.append(pdf)

        merger.write(output_file)
        merger.close()

    """ Main function of EBookContentGenerator. Generates a full ebook using AI. Returns the title of the created ebook."""
    def generate_ebook(
        self,
        topic,
        target_audience,
        id,
        num_chapters=6,
        num_subsections=4,
        preview=True,
    ):
        if platform.system() == "Windows":
            import pythoncom

            pythoncom.CoInitialize()

        log_file = f"app/ai_book_generation/logs/log-{id}.txt"
        docx_file = f"app/ai_book_generation/docs/docs-{id}.docx"

        self.logger = open(log_file, "w")

        # Generate title and outline
        title = self.generate_title(topic, target_audience)

        t = Template(self.templates_file)
        # template = t.choose_template(title, topic, target_audience)
        template = {
            "cover_template": (
                "app/ai_book_generation/templates/covers/gen.docx"
            ),
            "book_template": (
                "app/ai_book_generation/templates/content/theme.docx"
            ),
        }

        self.generate_cover(
            template.get("cover_template"),
            title,
            topic,
            target_audience,
            f"{self.cover_pdf_location}",
            preview
        )

        outline = self.generate_outline(
            topic,
            target_audience,
            title,
            num_chapters,
            num_subsections,
        )

        self.generate_docx(
            topic,
            target_audience,
            title,
            outline,
            docx_file,
            template.get("book_template"),
            preview,
        )

        self.logger.close()

        # cover - cover.docx, book - docx_file
        # need to convert them both to pdfs and then merge them!
        convert(docx_file, self.content_pdf_location_temp)
        self.remove_first_page(
            self.content_pdf_location_temp, self.content_pdf_location
        )

        self.merge_pdfs(
            [self.cover_pdf_location, self.content_pdf_location],
            self.output_pdf_location,
        )

        page_count = self.get_ebook_page_count(self.output_pdf_location)

        if platform.system() == "Windows":
            pythoncom.CoUninitialize()

        return Ebook(
            title=title,
            topic=topic,
            target_audience=target_audience,
            docx_file=docx_file,
            pdf_file=self.output_pdf_location,
            cover_img=None,
            assets=None,
            page_count=page_count,
        )

    """ Function to get ebook page count """

    def get_ebook_page_count(self, pdf_file):
        with open(pdf_file, "rb") as file:
            # store data in pdfReader
            pdfReader = PyPDF2.PdfFileReader(file)
            page_count = pdfReader.numPages
            print(f"Total Pages: {page_count}")
            return page_count

    """ Function to remove the first page of PDF """

    def remove_first_page(self, source_pdf, output_pdf):
        # Open the source PDF file
        with open(source_pdf, "rb") as file:
            reader = PyPDF2.PdfFileReader(file)

            # Create a PDF writer object
            writer = PyPDF2.PdfFileWriter()

            # Number of pages in the source PDF
            num_pages = reader.numPages

            # Check if the PDF has more than one page
            if num_pages < 2:
                raise ValueError(
                    "The PDF has only one page and cannot be processed."
                )

            # Loop through all pages except the first one
            for page_num in range(1, num_pages):
                # Get the page
                page = reader.getPage(page_num)

                # Add it to the writer object
                writer.addPage(page)

            # Write out the new PDF
            with open(output_pdf, "wb") as output_file:
                writer.write(output_file)

    """Temporary test generate image function"""

    def generate_image(self):
        openai.OPENAI_API_KEY = self.OPENAI_API_KEY
        response = openai.Image.create(
            prompt="a white siamese cat", n=1, size="1024x1024"
        )
        image_url = response["data"][0]["url"]
        print(image_url)

    def generate_title(
        self,
        topic,
        target_audience,
    ):
        convo_id = self.GPT_WRAPPER.start_convo(GptSystems.AUTHOR_SYSTEM)
        # Generate title
        title_prompt = (
            f'We are writing an eBook. It is about "{topic}". Our'
            f' reader is:  "{target_audience}". Write a short, catch'
            " title clearly directed at our reader that is less than"
            " 9 words and proposes a “big promise” that will be sure to grab"
            " the readers attention."
        )
        title = self.GPT_WRAPPER.msg_in_convo(convo_id, title_prompt)
        # remove surrounding quotes from title (if any)
        title = title.replace('"', "")
        return title

    """
    Verify that ChatGPT is producing an outline in appropriate
    """

    def verify_outline(self, outline, num_chapters, num_subsections):
        # Check number of chapters
        if len(outline.items()) != num_chapters:
            return False
        for chapter, subtopics in outline.items():
            # Check that chapter includes chapter subject
            if len(chapter) < 15:
                return False
            # Check that subtopics are in a list
            if type(subtopics) != list:
                return False
            # Check number of subsections in each chapter
            if len(subtopics) != num_subsections:
                return False
        return True

    @retry(max_retries=3)
    def generate_outline(
        self,
        topic,
        target_audience,
        title,
        num_chapters,
        num_subsections,
    ):
        convo_id = self.GPT_WRAPPER.start_convo(GptSystems.AUTHOR_SYSTEM)
        outline_prompt = (
            f'We are writing an eBook called "{title}". It is about'
            f' "{topic}". Our reader is:  "{target_audience}".  Create'
            " a compehensive outline for our ebook, which will have"
            f" {num_chapters} chapter(s). Each chapter should have exactly"
            f" {num_subsections} subsection(s) Output Format for prompt:"
            " python dict with key: chapter title, value: a single list/array"
            " containing subsection titles within the chapter (the subtopics"
            " should be inside the list). The chapter titles should be"
            ' prepended with the chapter number, like this: "Chapter 5:'
            ' [chapter title]". The subsection titles should be prepended'
            ' with the {chapter number.subtitle number}, like this: "5.4:'
            ' [subsection title]". '
        )
        outline_json = self.GPT_WRAPPER.msg_in_convo(convo_id, outline_prompt)
        outline_json = outline_json[
            outline_json.find("{") : outline_json.rfind("}") + 1
        ]
        outline = json.loads(outline_json)
        if not self.verify_outline(outline, num_chapters, num_subsections):
            raise Exception("Outline not well formed!")

        return outline

    @retry(max_retries=1)
    def generate_chapter_content(
        self, title, topic, target_audience, idx, chapter, subtopic
    ):
        convo_id = self.GPT_WRAPPER.start_convo(GptSystems.AUTHOR_SYSTEM)
        num_words_str = "500 to 700"
        content_prompt = (
            f'We are writing an eBook called "{title}". Overall, it is about'
            f' "{topic}". Our reader is:  "{target_audience}". We are'
            f" currently writing the #{idx+1} section for the chapter:"
            f' "{chapter}". Using at least {num_words_str} words, write the'
            " full contents of the section regarding this subtopic:"
            f' "{subtopic}". The output should be as helpful to the reader as'
            " possible. Include quantitative facts and statistics, with"
            " references. Go as in depth as necessary. You can split this"
            " into multiple paragraphs if you see fit. The output should also"
            ' be in cohesive paragraph form. Do not include any "[Insert'
            ' ___]" parts that will require manual editing in the book later.'
            " If find yourself needing to put 'insert [blank]' anywhere, do"
            " not do it (this is very important). If you do not know"
            " something, do not include it in the output. Exclude any"
            " auxillary information like  the word count, as the entire"
            " output will go directly into the ebook for readers, without any"
            " human procesing. Remember the {num_words_str} word minimum,"
            " please adhere to it."
        )
        content = self.GPT_WRAPPER.msg_in_convo(convo_id, content_prompt)

        def remove_subtopic_from_content(content, subtopic, search_limit=200):
            # Define the pattern to search for X.X where X is any number
            pattern = r'\d\.\d'

            # Search for the pattern in the first 100 characters of the input string
            match = re.search(pattern, content[:search_limit])

            if match:
                # Find the index of the next newline character after the match
                newline_index = content.find('\n', match.end())

                if newline_index != -1:
                    # Extract the substring after the newline character
                    result_string = content[newline_index + 1:]
                    content = content.strip()
                    return result_string
            return content

        try:
            content = content.strip()
            content = remove_subtopic_from_content(content, subtopic)
        except:
            pass

        return content

    @retry(max_retries=1)
    def generate_actionable_steps(
        self,
        title,
        topic,
        target_audience,
        idx,
        chapter,
        subtopics,
        subtopics_content,
    ):
        # Generate actionable items for chapter
        content_prompt = (
            f'We are writing an eBook called "{title}". Overall, it is'
            f' about "{topic}". Our reader is: '
            f' "{target_audience}". We are currently writing the'
            f' #{idx+1} section for the chapter: "{chapter}". Write the'
            f" contents for the following subtopics: {subtopics}. The"
            " output should be as helpful to the reader as"
            " possible. Dive deep in the specifics as much as possible."
        )
        actions_prompt = (
            f"Based on the contents of this chapter, create a list of the"
            f" 3-4 most important actionable steps for the reader Keep"
            f" this as clear and consise as possible, sort of like a"
            f" checklist of the reader"
        )

        convo_id = self.GPT_WRAPPER.start_convo(GptSystems.AUTHOR_SYSTEM)
        content = self.GPT_WRAPPER.msg_in_convo_given_history(
            convo_id,
            messages=[
                {"role": "system", "content": GptSystems.AUTHOR_SYSTEM},
                {"role": "user", "content": content_prompt},
                {
                    "role": "assistant",
                    "content": " ".join(subtopics_content),
                },
                {"role": "user", "content": actions_prompt},
            ],
        )
        return content

    

    def char_can_be_encoded(self, char, encoding):
        try:
            char.encode(encoding)
            return True
        except UnicodeEncodeError:
            return False

    def remove_non_charmap(self, s, encoding='charmap'):
        return ''.join(c for c in s if self.char_can_be_encoded(c,encoding))

    @retry(max_retries=3)
    def generate_docx(
        self,
        topic,
        target_audience,
        title,
        outline,
        docx_file,
        book_template,
        preview,
        actionable_steps=False,
    ):
        document = Document(book_template)
        # style.paragraph_format.line_spacing = 1.15
        # Table of Contents is at this part of the document (from the template)
        # TODO: Add page numbers, takes quite a bit of work, might not be worth it right now
        document.add_page_break()  # Lets add a page and then remove it later. This is due to formatting issues within Word - let's just start with a fresh page
        document.add_heading("Table of Contents")
        for chapter, subtopics in outline.items():
            chapter = self.remove_non_charmap(chapter)
            document.add_heading(chapter, level=2)
            for idx, subtopic in enumerate(subtopics):
                subtopic = self.remove_non_charmap(subtopic)
                document.add_heading("\t" + subtopic, level=3)

        document.add_page_break()

        chapter_num = 1
        for chapter, subtopics in outline.items():
            chapter = self.remove_non_charmap(chapter)
            document.add_heading(chapter, level=1)
            subtopics_content = []

            # Generate each subtopic content
            for idx, subtopic in enumerate(subtopics):
                # Stop writing the ebook after four subsections if preview
                if preview and idx >= 2:
                    break

                subtopic = self.remove_non_charmap(subtopic)
                document.add_heading(subtopic, level=2)

                content = self.generate_chapter_content(
                    title, topic, target_audience, idx, chapter, subtopic
                )
                content = self.remove_non_charmap(content)
                document.add_paragraph(content)
                subtopics_content.append(content)

            # Stop writing the ebook after one chapter if a preview
            if preview:
                document.add_heading(
                    "Preview Completed - Purchase Full Book To Read More!"
                )
                break

            if actionable_steps:
                content = self.generate_actionable_steps(
                    title,
                    topic,
                    target_audience,
                    idx,
                    chapter,
                    subtopics,
                    subtopics_content,
                )
                document.add_heading("Actionable Steps", level=2)
                content = self.remove_non_charmap(content)
                document.add_paragraph(content)

            if chapter_num < len(outline.items()):
                document.add_page_break()
            chapter_num += 1

        document.save(docx_file)


if __name__ == "__main__":
    eg = EBookContentGenerator(
        templates_file="app/ai_book_generation/templates/templates.json",
        output_directory="app/ai_book_generation/output",
        id=2,
    )
    ebook = eg.generate_ebook(
        "guide to looksmaxxing",
        "mid-twenties coders",
        2,
        6,
        4,
        True,
    )
